#!/usr/bin/env python
"""
Resume Generator Script
Generates Word and PDF versions of professional resume from website content
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import os

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_word_resume():
    """Create the Word document resume"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Profile Picture
    profile_pic_path = r'D:\repos\dotnet\resume-site\myprofile.png'
    if os.path.exists(profile_pic_path):
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_para.add_run()
        pic_run.add_picture(profile_pic_path, width=Inches(1.2))
        pic_para.paragraph_format.space_after = Pt(6)

    # Header Section - Name and Title
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run('GANESAN SANKARAN')
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.name = 'Arial'

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('Software Principal Technical Expert')
    title_run.font.size = Pt(12)
    title_run.font.name = 'Arial'

    # Contact Information
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(
        'Hosur, Tamil Nadu, India | +91 95006 62498 | ganesan.ksr88@outlook.com | linkedin.com/in/ganesansnkr'
    )
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Arial'

    doc.add_paragraph()  # Spacing

    # Professional Summary
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
    summary_run.font.size = Pt(12)
    summary_run.font.bold = True
    summary_run.font.name = 'Arial'
    add_horizontal_line(summary_heading)

    summary = doc.add_paragraph(
        'Software Principal Technical Expert with 14+ years of experience in modern software development, '
        'cloud architecture, and team leadership. Expertise in .NET Core, C#, Azure, Docker, and Angular. '
        'Proven track record in leading cross-functional teams, implementing AI/ML solutions, and delivering '
        'scalable SaaS products with focus on code quality, security, and best practices.',
        style='Normal'
    )
    summary.paragraph_format.space_after = Pt(6)
    for run in summary.runs:
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    # Core Technical Skills
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run('CORE TECHNICAL SKILLS')
    skills_run.font.size = Pt(12)
    skills_run.font.bold = True
    skills_run.font.name = 'Arial'
    add_horizontal_line(skills_heading)

    skills_data = [
        ('Backend Development:', 'C#, .NET Core, ASP.NET, Web API, Entity Framework, Microservices, REST APIs, gRPC'),
        ('Frontend Development:', 'Angular, TypeScript, JavaScript, HTML5/CSS3, WPF, XAML, MVVM'),
        ('Cloud & DevOps:', 'Azure (App Service, Functions, Storage, Monitor), Docker, Kubernetes, CI/CD, Azure DevOps'),
        ('Databases:', 'SQL Server, PostgreSQL, Redis, MongoDB, Cosmos DB'),
        ('AI & Integration:', 'ML.NET, OpenAI, LLM, RabbitMQ, Azure Service Bus, Event-Driven Architecture'),
        ('Tools & Practices:', 'Git, SonarQube, HP Fortify, Agile/Scrum, Code Review, Unit Testing, Security')
    ]

    for label, skills in skills_data:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        label_run = p.add_run(label + ' ')
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        label_run.font.name = 'Arial'
        skills_run = p.add_run(skills)
        skills_run.font.size = Pt(10)
        skills_run.font.name = 'Arial'

    doc.add_paragraph()  # Spacing

    # Professional Experience
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
    exp_run.font.size = Pt(12)
    exp_run.font.bold = True
    exp_run.font.name = 'Arial'
    add_horizontal_line(exp_heading)

    # Experience 1: Schneider Electric
    exp1_title = doc.add_paragraph()
    exp1_title.paragraph_format.space_after = Pt(2)
    title_run = exp1_title.add_run('Software Principal Technical Expert')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp1_company = doc.add_paragraph()
    exp1_company.paragraph_format.space_after = Pt(4)
    company_run = exp1_company.add_run('Schneider Electric | Bengaluru, India | February 2022 – Present')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp1_bullets = [
        'Developed key SaaS energy management product using .NET 8, C#, Azure, and Angular for big buildings smart energy management',
        'Integrated AI/ML features using ML.NET and OpenAI for data-driven energy optimization and predictive analytics',
        'Led migration from Windows-based to cloud-based multi-tenant architecture, achieving zero-downtime deployment',
        'Improved maintenance efficiency by 50% through implementation of code quality gates using SonarQube and HP Fortify',
        'Containerized application services with Docker and implemented end-to-end monitoring using Azure Monitor and OpenTelemetry'
    ]
    for bullet in exp1_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Experience 2: Philips India
    exp2_title = doc.add_paragraph()
    exp2_title.paragraph_format.space_before = Pt(6)
    exp2_title.paragraph_format.space_after = Pt(2)
    title_run = exp2_title.add_run('Software Technical Lead')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp2_company = doc.add_paragraph()
    exp2_company.paragraph_format.space_after = Pt(4)
    company_run = exp2_company.add_run('Philips India LTD | Bengaluru, India | November 2018 – February 2022')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp2_bullets = [
        'Led team of 6 engineers in core R&D team for healthcare platform development using .NET Core, C#, and Angular',
        'Improved code quality and security by 40% through implementation of SonarQube, Black Duck, and HP Fortify scanning',
        'Implemented SDLC best practices and security standards, boosting code performance and reducing vulnerabilities',
        'Conducted weekly technical huddles on emerging technologies, improving team skill sets and knowledge sharing',
        'Collaborated with product owners on feature prioritization and technology advancement roadmaps'
    ]
    for bullet in exp2_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Experience 3: Aricent Technologies
    exp3_title = doc.add_paragraph()
    exp3_title.paragraph_format.space_before = Pt(6)
    exp3_title.paragraph_format.space_after = Pt(2)
    title_run = exp3_title.add_run('Technical Lead')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp3_company = doc.add_paragraph()
    exp3_company.paragraph_format.space_after = Pt(4)
    company_run = exp3_company.add_run('Aricent Technologies | Bengaluru, India | March 2015 – November 2018')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp3_bullets = [
        'Led team of 8 developers in end-to-end enterprise application development using ASP.NET, Azure, and Angular',
        'Managed continuous delivery cycles with Jenkins, achieving bi-weekly deployments with 99.5% success rate',
        'Conducted requirement reviews and stakeholder feedback sessions, ensuring alignment with business objectives',
        'Monitored application stability, security, and scalability post-release using Azure Monitor and custom dashboards'
    ]
    for bullet in exp3_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Add page break for page 2
    doc.add_page_break()

    # Experience 4-6: Earlier roles (condensed)
    exp4_title = doc.add_paragraph()
    exp4_title.paragraph_format.space_after = Pt(2)
    title_run = exp4_title.add_run('Senior Software Engineer')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp4_company = doc.add_paragraph()
    exp4_company.paragraph_format.space_after = Pt(4)
    company_run = exp4_company.add_run('Tagit Mobile Pvt Ltd | Chennai, India | June 2014 – March 2015')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp4_bullets = [
        'Developed custom controls for Windows Phone and tablet devices using C# and XAML',
        'Improved app performance and stability by 30% through memory leak fixes and optimization'
    ]
    for bullet in exp4_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Experience 5
    exp5_title = doc.add_paragraph()
    exp5_title.paragraph_format.space_before = Pt(6)
    exp5_title.paragraph_format.space_after = Pt(2)
    title_run = exp5_title.add_run('Senior Software Engineer')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp5_company = doc.add_paragraph()
    exp5_company.paragraph_format.space_after = Pt(4)
    company_run = exp5_company.add_run('Photon Interactive | Chennai, India | January 2014 – May 2014')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp5_bullets = [
        'Developed Windows Phone and Surface tablet apps for Fortune 500 clients (Walgreens, Best Buy, Chase Mobile)',
        'Created reusable custom controls, reducing development time by 25% across multiple projects'
    ]
    for bullet in exp5_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Experience 6
    exp6_title = doc.add_paragraph()
    exp6_title.paragraph_format.space_before = Pt(6)
    exp6_title.paragraph_format.space_after = Pt(2)
    title_run = exp6_title.add_run('Software Developer')
    title_run.font.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Arial'

    exp6_company = doc.add_paragraph()
    exp6_company.paragraph_format.space_after = Pt(4)
    company_run = exp6_company.add_run('Syncfusion Software Solutions | Chennai, India | November 2010 – January 2014')
    company_run.font.size = Pt(10)
    company_run.font.italic = True
    company_run.font.name = 'Arial'

    exp6_bullets = [
        'Developed data visualization components (charts, grids, schedulers) using WPF, Silverlight, and C#',
        'Enabled real-time data visualization for web and mobile platforms, serving 1000+ enterprise customers'
    ]
    for bullet in exp6_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    doc.add_paragraph()  # Spacing

    # Key Achievements
    achievements_heading = doc.add_paragraph()
    achievements_run = achievements_heading.add_run('KEY ACHIEVEMENTS')
    achievements_run.font.size = Pt(12)
    achievements_run.font.bold = True
    achievements_run.font.name = 'Arial'
    add_horizontal_line(achievements_heading)

    achievements = [
        'Improved maintenance efficiency by 50% through quality gates and code analysis tools implementation',
        'Successfully migrated enterprise applications to cloud with zero-downtime deployment',
        'Led teams of up to 8 engineers, mentoring junior developers and conducting code reviews',
        'Reduced development time by 25% through creation of reusable component libraries',
        'Achieved 99.5% deployment success rate with bi-weekly continuous delivery cycles'
    ]
    for achievement in achievements:
        p = doc.add_paragraph(achievement, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Key Projects
    projects_heading = doc.add_paragraph()
    projects_run = projects_heading.add_run('KEY PROJECTS')
    projects_run.font.size = Pt(12)
    projects_run.font.bold = True
    projects_run.font.name = 'Arial'
    add_horizontal_line(projects_heading)

    projects = [
        'Energy Management SaaS Platform: Multi-tenant cloud platform with AI/ML integration for big buildings smart energy management',
        'Healthcare R&D Platform: Core healthcare platform with enhanced security and SDLC best practices',
        'Enterprise Application Platform: Scalable application serving 1000+ users with bi-weekly deployments'
    ]
    for project in projects:
        p = doc.add_paragraph(project, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Education
    education_heading = doc.add_paragraph()
    education_run = education_heading.add_run('EDUCATION')
    education_run.font.size = Pt(12)
    education_run.font.bold = True
    education_run.font.name = 'Arial'
    add_horizontal_line(education_heading)

    edu = doc.add_paragraph()
    edu.paragraph_format.space_after = Pt(2)
    edu_degree = edu.add_run('Bachelor of Engineering (B.E.) - Computer Science and Engineering')
    edu_degree.font.bold = True
    edu_degree.font.size = Pt(10)
    edu_degree.font.name = 'Arial'

    edu_details = doc.add_paragraph()
    edu_details.paragraph_format.left_indent = Inches(0.25)
    edu_details_run = edu_details.add_run('Anna University Chennai, AKCE College of Engineering, Srivilliputhur, Tamilnadu | June 2006 – May 2010')
    edu_details_run.font.size = Pt(10)
    edu_details_run.font.name = 'Arial'

    # Certifications
    cert_heading = doc.add_paragraph()
    cert_run = cert_heading.add_run('CERTIFICATIONS')
    cert_run.font.size = Pt(12)
    cert_run.font.bold = True
    cert_run.font.name = 'Arial'
    add_horizontal_line(cert_heading)

    certifications = [
        'Certified Azure Developer (Microsoft Azure) - June 2023',
        'Certified SAFe 4 DevOps Practitioner (Scaled Agile) - July 2019'
    ]
    for cert in certifications:
        p = doc.add_paragraph(cert, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Languages
    lang_heading = doc.add_paragraph()
    lang_run = lang_heading.add_run('LANGUAGES')
    lang_run.font.size = Pt(12)
    lang_run.font.bold = True
    lang_run.font.name = 'Arial'
    add_horizontal_line(lang_heading)

    lang = doc.add_paragraph()
    lang.paragraph_format.left_indent = Inches(0.25)
    lang_run = lang.add_run('English (Fluent), Tamil (Native), German (Basic)')
    lang_run.font.size = Pt(10)
    lang_run.font.name = 'Arial'

    # Save Word document
    output_path = r'D:\repos\dotnet\resume-site\assets\documents\Ganesan_Sankaran.docx'
    doc.save(output_path)
    print(f"[OK] Word document created: {output_path}")
    return output_path

def create_pdf_resume():
    """Create the PDF resume using ReportLab"""
    output_path = r'D:\repos\dotnet\resume-site\assets\documents\Ganesan_Sankaran.pdf'

    # Create document
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch
    )

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='black',
        alignment=TA_CENTER,
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )

    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=12,
        alignment=TA_CENTER,
        spaceAfter=6,
        fontName='Helvetica'
    )

    contact_style = ParagraphStyle(
        'Contact',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName='Helvetica'
    )

    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='black',
        spaceAfter=6,
        spaceBefore=6,
        fontName='Helvetica-Bold',
        borderWidth=1,
        borderColor='black',
        borderPadding=2,
        borderRadius=0
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6,
        fontName='Helvetica'
    )

    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceAfter=3,
        fontName='Helvetica',
        bulletIndent=10
    )

    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceAfter=2
    )

    company_style = ParagraphStyle(
        'Company',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica-Oblique',
        spaceAfter=4
    )

    # Header
    # Profile Picture
    profile_pic_path = r'D:\repos\dotnet\resume-site\myprofile.png'
    if os.path.exists(profile_pic_path):
        img = Image(profile_pic_path, width=1.2*inch, height=1.2*inch)
        img.hAlign = 'CENTER'
        elements.append(img)
        elements.append(Spacer(1, 0.1*inch))

    elements.append(Paragraph('GANESAN SANKARAN', title_style))
    elements.append(Paragraph('Software Principal Technical Expert', subtitle_style))
    elements.append(Paragraph(
        'Hosur, Tamil Nadu, India | +91 95006 62498 | ganesan.ksr88@outlook.com | linkedin.com/in/ganesansnkr',
        contact_style
    ))

    # Professional Summary
    elements.append(Paragraph('<b>PROFESSIONAL SUMMARY</b>', heading_style))
    elements.append(Paragraph(
        'Software Principal Technical Expert with 14+ years of experience in modern software development, '
        'cloud architecture, and team leadership. Expertise in .NET Core, C#, Azure, Docker, and Angular. '
        'Proven track record in leading cross-functional teams, implementing AI/ML solutions, and delivering '
        'scalable SaaS products with focus on code quality, security, and best practices.',
        normal_style
    ))

    # Core Technical Skills
    elements.append(Paragraph('<b>CORE TECHNICAL SKILLS</b>', heading_style))
    skills_data = [
        ('<b>Backend Development:</b> C#, .NET Core, ASP.NET, Web API, Entity Framework, Microservices, REST APIs, gRPC'),
        ('<b>Frontend Development:</b> Angular, TypeScript, JavaScript, HTML5/CSS3, WPF, XAML, MVVM'),
        ('<b>Cloud &amp; DevOps:</b> Azure (App Service, Functions, Storage, Monitor), Docker, Kubernetes, CI/CD, Azure DevOps'),
        ('<b>Databases:</b> SQL Server, PostgreSQL, Redis, MongoDB, Cosmos DB'),
        ('<b>AI &amp; Integration:</b> ML.NET, OpenAI, LLM, RabbitMQ, Azure Service Bus, Event-Driven Architecture'),
        ('<b>Tools &amp; Practices:</b> Git, SonarQube, HP Fortify, Agile/Scrum, Code Review, Unit Testing, Security')
    ]
    for skill in skills_data:
        elements.append(Paragraph(skill, normal_style))

    elements.append(Spacer(1, 0.1*inch))

    # Professional Experience
    elements.append(Paragraph('<b>PROFESSIONAL EXPERIENCE</b>', heading_style))

    # Schneider Electric
    elements.append(Paragraph('Software Principal Technical Expert', job_title_style))
    elements.append(Paragraph('Schneider Electric | Bengaluru, India | February 2022 – Present', company_style))
    exp1_bullets = [
        'Developed key SaaS energy management product using .NET 8, C#, Azure, and Angular for big buildings smart energy management',
        'Integrated AI/ML features using ML.NET and OpenAI for data-driven energy optimization and predictive analytics',
        'Led migration from Windows-based to cloud-based multi-tenant architecture, achieving zero-downtime deployment',
        'Improved maintenance efficiency by 50% through implementation of code quality gates using SonarQube and HP Fortify',
        'Containerized application services with Docker and implemented end-to-end monitoring using Azure Monitor'
    ]
    for bullet in exp1_bullets:
        elements.append(Paragraph(f'• {bullet}', bullet_style))

    # Philips India
    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph('Software Technical Lead', job_title_style))
    elements.append(Paragraph('Philips India LTD | Bengaluru, India | November 2018 – February 2022', company_style))
    exp2_bullets = [
        'Led team of 6 engineers in core R&D team for healthcare platform development using .NET Core, C#, and Angular',
        'Improved code quality and security by 40% through implementation of SonarQube, Black Duck, and HP Fortify',
        'Implemented SDLC best practices and security standards, boosting code performance and reducing vulnerabilities',
        'Conducted weekly technical huddles on emerging technologies, improving team skill sets',
        'Collaborated with product owners on feature prioritization and technology advancement roadmaps'
    ]
    for bullet in exp2_bullets:
        elements.append(Paragraph(f'• {bullet}', bullet_style))

    # Aricent
    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph('Technical Lead', job_title_style))
    elements.append(Paragraph('Aricent Technologies | Bengaluru, India | March 2015 – November 2018', company_style))
    exp3_bullets = [
        'Led team of 8 developers in end-to-end enterprise application development using ASP.NET, Azure, and Angular',
        'Managed continuous delivery with Jenkins, achieving bi-weekly deployments with 99.5% success rate',
        'Monitored application stability, security, and scalability using Azure Monitor and custom dashboards'
    ]
    for bullet in exp3_bullets:
        elements.append(Paragraph(f'• {bullet}', bullet_style))

    # Page break
    elements.append(PageBreak())

    # Earlier roles
    elements.append(Paragraph('Senior Software Engineer', job_title_style))
    elements.append(Paragraph('Tagit Mobile Pvt Ltd | Chennai, India | June 2014 – March 2015', company_style))
    elements.append(Paragraph('• Developed custom controls for Windows Phone and tablet devices using C# and XAML', bullet_style))
    elements.append(Paragraph('• Improved app performance and stability by 30% through memory leak fixes', bullet_style))

    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph('Senior Software Engineer', job_title_style))
    elements.append(Paragraph('Photon Interactive | Chennai, India | January 2014 – May 2014', company_style))
    elements.append(Paragraph('• Developed Windows Phone and Surface apps for Fortune 500 clients (Walgreens, Best Buy, Chase)', bullet_style))
    elements.append(Paragraph('• Created reusable custom controls, reducing development time by 25%', bullet_style))

    elements.append(Spacer(1, 0.05*inch))
    elements.append(Paragraph('Software Developer', job_title_style))
    elements.append(Paragraph('Syncfusion Software Solutions | Chennai, India | November 2010 – January 2014', company_style))
    elements.append(Paragraph('• Developed data visualization components using WPF, Silverlight, and C#', bullet_style))
    elements.append(Paragraph('• Enabled real-time data visualization for 1000+ enterprise customers', bullet_style))

    elements.append(Spacer(1, 0.1*inch))

    # Key Achievements
    elements.append(Paragraph('<b>KEY ACHIEVEMENTS</b>', heading_style))
    achievements = [
        'Improved maintenance efficiency by 50% through quality gates implementation',
        'Successfully migrated enterprise applications to cloud with zero-downtime deployment',
        'Led teams of up to 8 engineers, mentoring junior developers and conducting code reviews',
        'Reduced development time by 25% through reusable component libraries'
    ]
    for achievement in achievements:
        elements.append(Paragraph(f'• {achievement}', bullet_style))

    # Key Projects
    elements.append(Paragraph('<b>KEY PROJECTS</b>', heading_style))
    projects = [
        'Energy Management SaaS Platform: Multi-tenant cloud platform with AI/ML integration for big buildings smart energy management',
        'Healthcare R&D Platform: Core platform with enhanced security and SDLC best practices',
        'Enterprise Application Platform: Scalable application serving 1000+ users'
    ]
    for project in projects:
        elements.append(Paragraph(f'• {project}', bullet_style))

    # Education
    elements.append(Paragraph('<b>EDUCATION</b>', heading_style))
    elements.append(Paragraph('<b>Bachelor of Engineering (B.E.) - Computer Science and Engineering</b>', normal_style))
    elements.append(Paragraph('Anna University Chennai, AKCE College of Engineering, Srivilliputhur, Tamilnadu | June 2006 – May 2010', bullet_style))

    # Certifications
    elements.append(Paragraph('<b>CERTIFICATIONS</b>', heading_style))
    elements.append(Paragraph('• Certified Azure Developer (Microsoft Azure) - June 2023', bullet_style))
    elements.append(Paragraph('• Certified SAFe 4 DevOps Practitioner (Scaled Agile) - July 2019', bullet_style))

    # Languages
    elements.append(Paragraph('<b>LANGUAGES</b>', heading_style))
    elements.append(Paragraph('English (Fluent), Tamil (Native), German (Basic)', normal_style))

    # Build PDF
    doc.build(elements)
    print(f"[OK] PDF document created: {output_path}")
    return output_path

if __name__ == '__main__':
    print("Starting resume generation...")
    print("-" * 50)

    try:
        # Generate Word document
        word_path = create_word_resume()

        # Generate PDF document
        pdf_path = create_pdf_resume()

        print("-" * 50)
        print("[OK] Resume generation completed successfully!")
        print(f"\nGenerated files:")
        print(f"  - Word: {word_path}")
        print(f"  - PDF:  {pdf_path}")

    except Exception as e:
        print(f"[ERROR] Error generating resume: {str(e)}")
        import traceback
        traceback.print_exc()
